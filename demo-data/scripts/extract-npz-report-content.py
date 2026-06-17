from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document


KEY_TAKEAWAYS_PREFIX = "Ключевые выводы:"
GOAL_PREFIX = "Цель:"
MATERIAL_PREFIX = "Учебный материал:"
VIDEO_MARKER = "Видео:"
TEST_MARKER = "Тест:"
TEST_NAME_PREFIX = "Название теста:"
TECH_SECTION_PREFIX = "Технический раздел:"


def clean_text(value: str) -> str:
    return value.replace("\xa0", " ").replace("\u2011", "-").strip()


def strip_prefix(value: str, prefix: str) -> str:
    return value[len(prefix) :].strip() if value.startswith(prefix) else value.strip()


def parse_options(raw_options: str, raw_correct: str, topic_code: str, question_number: int) -> list[dict]:
    option_matches = re.findall(r"([A-D])\)\s*(.*?)(?=(?:;\s*[A-D]\)|$))", raw_options)
    if not option_matches:
        raise ValueError(
            f"Cannot parse answer options for {topic_code} question {question_number}: {raw_options!r}"
        )

    correct_letters = set(re.findall(r"[A-D]", raw_correct))
    return [
        {
            "code": f"OPT-{topic_code[6:]}-{question_number:02d}-{letter}",
            "body": body.strip(),
            "answerOptionRole": "CHOICE_OPTION",
            "isCorrect": letter in correct_letters,
            "displayOrder": display_order,
        }
        for display_order, (letter, body) in enumerate(option_matches)
    ]


def build_material_body(body_paragraphs: list[str], key_takeaways: str, videos: list[dict]) -> str:
    sections: list[str] = []
    sections.extend(body_paragraphs)
    sections.append(f"Ключевые выводы: {key_takeaways}")
    if videos:
        video_lines = ["Рекомендуемые видео:"]
        for video in videos:
            video_lines.append(f"- {video['title']}")
            video_lines.append(f"  Ссылка: {video['url']}")
            video_lines.append(f"  Источник: {video['source']}")
            video_lines.append(f"  Почему подходит: {video['why']}")
        sections.append("\n".join(video_lines))
    return "\n\n".join(section for section in sections if section.strip())


def parse_report(docx_path: Path) -> dict:
    document = Document(docx_path)
    paragraphs = [clean_text(paragraph.text) for paragraph in document.paragraphs]
    paragraphs = [paragraph for paragraph in paragraphs if paragraph]
    tables = [
        [[clean_text(cell.text).replace("\n", " | ") for cell in row.cells] for row in table.rows]
        for table in document.tables
    ]

    courses: list[dict] = []
    current_course: dict | None = None
    # The first table is the high-level course summary and does not belong to a topic block.
    table_index = 1
    index = 0

    while index < len(paragraphs):
        line = paragraphs[index]

        if line.startswith("COURSE-"):
            course_code, course_name = [part.strip() for part in line.split(".", 1)]
            current_course = {
                "code": course_code,
                "name": course_name,
                "description": paragraphs[index + 2],
                "sortOrder": len(courses),
                "topics": [],
            }
            courses.append(current_course)
            index += 4
            continue

        if line.startswith("TOPIC-"):
            if current_course is None:
                raise ValueError(f"Topic found before course: {line}")

            topic_code, topic_name = [part.strip() for part in line.split(".", 1)]
            topic_description = strip_prefix(paragraphs[index + 1], GOAL_PREFIX)

            material_paragraphs: list[str] = []
            cursor = index + 2
            while cursor < len(paragraphs) and not paragraphs[cursor].startswith(KEY_TAKEAWAYS_PREFIX):
                material_paragraphs.append(strip_prefix(paragraphs[cursor], MATERIAL_PREFIX))
                cursor += 1

            if cursor >= len(paragraphs):
                raise ValueError(f"Missing key takeaways block for topic {topic_code}")

            key_takeaways = strip_prefix(paragraphs[cursor], KEY_TAKEAWAYS_PREFIX)

            if paragraphs[cursor + 1] != VIDEO_MARKER or paragraphs[cursor + 2] != TEST_MARKER:
                raise ValueError(
                    f"Unexpected topic markers for {topic_code}: "
                    f"{paragraphs[cursor + 1]!r}, {paragraphs[cursor + 2]!r}"
                )

            test_name = strip_prefix(paragraphs[cursor + 3], TEST_NAME_PREFIX).rstrip(".").strip("«»")

            if table_index + 1 >= len(tables):
                raise ValueError(f"Not enough tables left to parse topic {topic_code}")

            video_rows = tables[table_index]
            question_rows = tables[table_index + 1]
            table_index += 2

            videos = [
                {
                    "title": row[0],
                    "url": row[1],
                    "source": row[2],
                    "why": row[3],
                }
                for row in video_rows[1:]
            ]

            questions = []
            for row in question_rows[1:]:
                question_number = int(row[0])
                questions.append(
                    {
                        "code": f"Q-{topic_code[6:]}-{question_number:02d}",
                        "body": row[2],
                        "questionType": "SINGLE_CHOICE" if "Single" in row[1] else "MULTIPLE_CHOICE",
                        "sortOrder": question_number - 1,
                        "explanation": row[5],
                        "options": parse_options(row[3], row[4], topic_code, question_number),
                    }
                )

            current_course["topics"].append(
                {
                    "code": topic_code,
                    "name": topic_name,
                    "description": topic_description,
                    "sortOrder": len(current_course["topics"]),
                    "material": {
                        "code": f"MAT-{topic_code[6:]}",
                        "name": topic_name,
                        "description": topic_description,
                        "body": build_material_body(material_paragraphs, key_takeaways, videos),
                        "materialType": "TEXT",
                        "sortOrder": 0,
                        "videoReferences": videos,
                    },
                    "test": {
                        "code": f"TEST-{topic_code[6:]}",
                        "name": test_name,
                        "description": f"Тест по теме: {topic_name}",
                        "testType": "CONTROL",
                        "thresholdPercent": 80,
                        "scoringPolicyCode": "DEFAULT",
                        "sortOrder": 0,
                        "assignAsActiveFinal": True,
                        "questions": questions,
                    },
                }
            )

            index = cursor + 4
            continue

        if line.startswith(TECH_SECTION_PREFIX):
            break

        index += 1

    return {
        "sourceDocument": str(docx_path),
        "summary": {
            "courseCount": len(courses),
            "topicCount": sum(len(course["topics"]) for course in courses),
            "materialCount": sum(len(course["topics"]) for course in courses),
            "testCount": sum(len(course["topics"]) for course in courses),
            "questionCount": sum(
                len(topic["test"]["questions"])
                for course in courses
                for topic in course["topics"]
            ),
        },
        "courses": courses,
    }


def main(argv: list[str]) -> int:
    if len(argv) != 3:
        print(
            "Usage: extract-npz-report-content.py <input.docx> <output.json>",
            file=sys.stderr,
        )
        return 1

    input_path = Path(argv[1]).resolve()
    output_path = Path(argv[2]).resolve()
    dataset = parse_report(input_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(dataset, ensure_ascii=False, indent=2), encoding="utf-8")

    print(
        json.dumps(
            {
                "output": str(output_path),
                **dataset["summary"],
            },
            ensure_ascii=False,
            indent=2,
        )
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
